import streamlit as st
from PyPDF2 import PdfReader
import docx
from sentence_transformers import SentenceTransformer, util
from google import genai
from google.genai import types
import io

# -------------------- Config --------------------
API_KEY = "AIzaSyAd-pJNXnxFNCr-zrb967WAILMuD6O-QxQ"
st.set_page_config(
    page_title="Legal Document Summarizer", 
    layout="wide",
    page_icon="📄",
    initial_sidebar_state="collapsed"
)

# Load embedding model
@st.cache_resource
def load_embedding_model():
    return SentenceTransformer("all-MiniLM-L6-v2")

model = load_embedding_model()

# -------------------- Helper Functions --------------------
def extract_paragraphs(uploaded_file):
    """
    FIXED: Now takes the uploaded_file object directly, not just the name
    """
    paragraphs = []
    
    # Get file extension from the name
    file_extension = uploaded_file.name.lower().split('.')[-1]

    if file_extension == "txt":
        # Read bytes and decode
        text = uploaded_file.read().decode("utf-8")
        paragraphs = text.split("\n\n")

    elif file_extension == "pdf":
        reader = PdfReader(io.BytesIO(uploaded_file.read()))
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        paragraphs = text.split("\n\n")

    elif file_extension == "docx":
        doc = docx.Document(io.BytesIO(uploaded_file.read()))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Assign IDs
    indexed_paragraphs = [{"id": idx, "text": para.strip()} for idx, para in enumerate(paragraphs) if para.strip()]
    return indexed_paragraphs


def summarize_file(uploaded_file, model_name="gemini-2.5-flash"):
    """
    FIXED: Now takes uploaded_file object and handles it properly
    """
    # Reset file pointer to beginning (important!)
    uploaded_file.seek(0)
    file_bytes = uploaded_file.read()
    
    # Determine MIME type based on file extension
    file_extension = uploaded_file.name.lower().split('.')[-1]
    if file_extension == 'pdf':
        mime_type = 'application/pdf'
    elif file_extension == 'docx':
        mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    else:  # txt
        mime_type = 'text/plain'

    file_part = types.Part.from_bytes(
        data=file_bytes,
        mime_type=mime_type,
    )

    prompt = """
    Summarize the document strictly as plain bullet points.
    Rules:
    - Only return bullet points.
    - Do not include titles, labels, "insight", "priority", or any extra headers.
    - Each bullet should be short, clear, and factual.
    - No numbering, no categories, no metadata — just the bullet points.
    """

    client = genai.Client(api_key=API_KEY)
    response = client.models.generate_content(
        model=model_name,
        contents=[file_part, prompt]
    )

    # Split response into sentences/points
    summary_points = [p.strip() for p in response.text.split("\n") if p.strip()]
    return summary_points

def map_summary_to_paragraphs(paragraphs, summary_sentences):
    para_texts = [p["text"] for p in paragraphs]
    para_embeddings = model.encode(para_texts, convert_to_tensor=True)
    summary_embeddings = model.encode(summary_sentences, convert_to_tensor=True)

    mapping = []
    for idx, s_emb in enumerate(summary_embeddings):
        cos_scores = util.cos_sim(s_emb, para_embeddings)[0]
        best_match_idx = int(cos_scores.argmax())
        mapping.append({
            "summary_sentence": summary_sentences[idx],
            "reference_paragraph_id": paragraphs[best_match_idx]["id"],
            "reference_text": paragraphs[best_match_idx]["text"]
        })
    return mapping

# -------------------- Streamlit UI --------------------

# Add global CSS for the entire app
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');

* {
    font-family: 'Inter', sans-serif;
}

/* Hide Streamlit default elements */
#MainMenu {visibility: hidden;}
footer {visibility: hidden;}
header {visibility: hidden;}

/* Main app background */
.stApp {
    background: linear-gradient(135deg, #667eea 0%, #764ba2 50%, #f093fb 100%);
    min-height: 100vh;
}

/* Custom title styling */
.main-title {
    background: linear-gradient(135deg, #ffffff, #f8f9fa, #ffffff);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    font-size: 3.5rem;
    font-weight: 800;
    text-align: center;
    margin: 30px 0;
    text-shadow: 0 4px 20px rgba(255,255,255,0.3);
    animation: titleGlow 3s ease-in-out infinite alternate;
    filter: drop-shadow(2px 2px 4px rgba(0,0,0,0.3));
}

@keyframes titleGlow {
    0% { filter: drop-shadow(2px 2px 4px rgba(0,0,0,0.3)) brightness(1); }
    100% { filter: drop-shadow(2px 2px 4px rgba(0,0,0,0.3)) brightness(1.3); }
}
}

/* Upload container styling */
.upload-container {
    background: rgba(255, 255, 255, 0.95);
    backdrop-filter: blur(20px);
    border-radius: 25px;
    padding: 40px;
    margin: 30px auto;
    max-width: 800px;
    box-shadow: 0 25px 50px rgba(0,0,0,0.15);
    border: 1px solid rgba(255,255,255,0.2);
    position: relative;
    overflow: hidden;
}

.upload-container::before {
    content: '';
    position: absolute;
    top: 0;
    left: -100%;
    width: 100%;
    height: 4px;
    background: linear-gradient(90deg, #6c5ce7, #a29bfe, #fd79a8, #fdcb6e);
    animation: progressBar 3s infinite;
}

@keyframes progressBar {
    0% { left: -100%; }
    100% { left: 100%; }
}

/* Processing status styling */
.processing-card {
    background: rgba(255, 255, 255, 0.9);
    backdrop-filter: blur(15px);
    border-radius: 20px;
    padding: 30px;
    margin: 20px 0;
    box-shadow: 0 15px 35px rgba(0,0,0,0.1);
    border-left: 5px solid #6c5ce7;
    transition: all 0.4s ease;
}

.processing-card:hover {
    transform: translateY(-5px);
    box-shadow: 0 25px 50px rgba(108, 92, 231, 0.2);
}

/* File info styling */
.file-info {
    background: linear-gradient(135deg, #6c5ce7 0%, #a29bfe 100%);
    color: white;
    padding: 20px 30px;
    border-radius: 20px;
    margin: 20px 0;
    display: flex;
    align-items: center;
    box-shadow: 0 10px 30px rgba(108, 92, 231, 0.3);
    animation: slideInRight 0.6s ease-out;
}

@keyframes slideInRight {
    from {
        opacity: 0;
        transform: translateX(50px);
    }
    to {
        opacity: 1;
        transform: translateX(0);
    }
}

.file-info .file-icon {
    font-size: 2.5em;
    margin-right: 20px;
    animation: bounce 2s infinite;
}

@keyframes bounce {
    0%, 20%, 50%, 80%, 100% { transform: translateY(0); }
    40% { transform: translateY(-10px); }
    60% { transform: translateY(-5px); }
}

/* Success message styling */
.success-message {
    background: linear-gradient(135deg, #00b894, #00a085);
    color: white;
    padding: 20px;
    border-radius: 15px;
    margin: 15px 0;
    text-align: center;
    box-shadow: 0 8px 25px rgba(0, 184, 148, 0.3);
    animation: successPulse 0.6s ease-out;
}

@keyframes successPulse {
    0% { transform: scale(0.8); opacity: 0; }
    100% { transform: scale(1); opacity: 1; }
}

/* Generate button styling */
.stButton > button {
    background: linear-gradient(135deg, #6c5ce7 0%, #a29bfe 100%) !important;
    color: white !important;
    border: none !important;
    border-radius: 50px !important;
    padding: 15px 40px !important;
    font-size: 1.1rem !important;
    font-weight: 600 !important;
    text-transform: uppercase !important;
    letter-spacing: 1px !important;
    box-shadow: 0 10px 30px rgba(108, 92, 231, 0.4) !important;
    transition: all 0.4s cubic-bezier(0.4, 0.0, 0.2, 1) !important;
    position: relative !important;
    overflow: hidden !important;
    min-width: 200px !important;
}

.stButton > button:hover {
    transform: translateY(-3px) scale(1.05) !important;
    box-shadow: 0 20px 40px rgba(108, 92, 231, 0.6) !important;
    background: linear-gradient(135deg, #5a4fcf 0%, #8b7ed8 100%) !important;
}

.stButton > button:active {
    transform: translateY(-1px) scale(1.02) !important;
}

/* File uploader styling */
.stFileUploader > div > div {
    background: rgba(255, 255, 255, 0.95) !important;
    backdrop-filter: blur(20px) !important;
    border: 2px dashed #6c5ce7 !important;
    border-radius: 20px !important;
    padding: 40px 20px !important;
    text-align: center !important;
    transition: all 0.3s ease !important;
    box-shadow: 0 10px 30px rgba(108, 92, 231, 0.1) !important;
}

.stFileUploader > div > div:hover {
    border-color: #a29bfe !important;
    background: rgba(255, 255, 255, 1) !important;
    transform: translateY(-5px) !important;
    box-shadow: 0 20px 40px rgba(108, 92, 231, 0.2) !important;
}

.stFileUploader label {
    color: #2d3436 !important;
    font-weight: 600 !important;
    font-size: 1.1rem !important;
}

.stFileUploader small {
    color: #636e72 !important;
    font-size: 0.95rem !important;
}

/* Spinner styling */
.stSpinner > div {
    border-top-color: #6c5ce7 !important;
}

/* Processing steps styling */
.processing-step {
    display: flex;
    align-items: center;
    padding: 15px 0;
    border-bottom: 1px solid rgba(108, 92, 231, 0.1);
}

.processing-step:last-child {
    border-bottom: none;
}

.step-icon {
    background: linear-gradient(135deg, #6c5ce7, #a29bfe);
    color: white;
    width: 40px;
    height: 40px;
    border-radius: 50%;
    display: flex;
    align-items: center;
    justify-content: center;
    margin-right: 15px;
    font-weight: bold;
}

.step-text {
    flex: 1;
    font-size: 1.1rem;
    color: #2d3436;
}

/* Reset button styling */
.reset-button {
    background: linear-gradient(135deg, #fd79a8, #fdcb6e) !important;
    color: white !important;
    border: none !important;
    border-radius: 25px !important;
    padding: 10px 25px !important;
    font-weight: 500 !important;
    box-shadow: 0 5px 15px rgba(253, 121, 168, 0.3) !important;
    transition: all 0.3s ease !important;
}

.reset-button:hover {
    transform: translateY(-2px) !important;
    box-shadow: 0 8px 25px rgba(253, 121, 168, 0.5) !important;
}
</style>
""", unsafe_allow_html=True)

# Custom title with enhanced styling
st.markdown("""
<div class="main-title">
    📄 Legal Document Summarizer
</div>
<div style="text-align: center; color: rgba(255,255,255,0.9); font-size: 1.3rem; 
     margin-bottom: 40px; font-weight: 300;">
    AI-Powered Document Analysis with Intelligent Paragraph Mapping
</div>
""", unsafe_allow_html=True)

uploaded_file = st.file_uploader(
    "", 
    type=["pdf","docx","txt"],
    help="Upload your legal document for AI analysis"
)

# Wrap upload area in custom container
if not uploaded_file:
    st.markdown("""
    <div class="upload-container">
        <div style="text-align: center;">
            <div style="font-size: 4rem; margin-bottom: 20px; animation: float 3s ease-in-out infinite;">
                📁
            </div>
            <h2 style="color: #2d3436; margin-bottom: 15px; font-weight: 600;">
                Upload Your Legal Document
            </h2>
            <p style="color: #636e72; font-size: 1.1rem; margin-bottom: 30px;">
                Supports PDF, DOCX, and TXT files • Maximum 200MB
            </p>
            <div style="display: flex; justify-content: center; gap: 30px; margin-top: 30px;">
                <div style="text-align: center;">
                    <div style="font-size: 2rem; margin-bottom: 10px;">📄</div>
                    <div style="color: #636e72; font-size: 0.9rem;">PDF Files</div>
                </div>
                <div style="text-align: center;">
                    <div style="font-size: 2rem; margin-bottom: 10px;">📝</div>
                    <div style="color: #636e72; font-size: 0.9rem;">Word Docs</div>
                </div>
                <div style="text-align: center;">
                    <div style="font-size: 2rem; margin-bottom: 10px;">📋</div>
                    <div style="color: #636e72; font-size: 0.9rem;">Text Files</div>
                </div>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

if uploaded_file:
    # Initialize processing state
    if "processing_complete" not in st.session_state:
        st.session_state.processing_complete = False
    if "paragraphs" not in st.session_state:
        st.session_state.paragraphs = None
    if "mapping" not in st.session_state:
        st.session_state.mapping = None

    # Show processing interface only if not complete
    if not st.session_state.processing_complete:
        
        # Enhanced file info display
        st.markdown(f"""
        <div class="file-info">
            <div class="file-icon">📄</div>
            <div>
                <div style="font-size: 1.3rem; font-weight: 600; margin-bottom: 5px;">
                    {uploaded_file.name}
                </div>
                <div style="opacity: 0.9; font-size: 0.95rem;">
                    Size: {uploaded_file.size / 1024:.1f} KB • Ready for processing
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        # Processing steps preview
        st.markdown("""
        <div class="processing-card">
            <h3 style="color: #2d3436; margin-bottom: 25px; text-align: center;">
                🚀 Processing Pipeline
            </h3>
            <div class="processing-step">
                <div class="step-icon">1</div>
                <div class="step-text">Extract and analyze document paragraphs</div>
            </div>
            <div class="processing-step">
                <div class="step-icon">2</div>
                <div class="step-text">Generate AI-powered executive summary</div>
            </div>
            <div class="processing-step">
                <div class="step-icon">3</div>
                <div class="step-text">Map summary points to source paragraphs</div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        # Center the generate button
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            if st.button("🚀 Start Analysis", type="primary"):
                # Step 1: Extract paragraphs
                with st.spinner("🔍 Extracting paragraphs..."):
                    paragraphs = extract_paragraphs(uploaded_file)
                
                st.markdown(f"""
                <div class="success-message">
                    ✅ Successfully extracted {len(paragraphs)} paragraphs
                </div>
                """, unsafe_allow_html=True)

                # Step 2: Generate summary
                with st.spinner("🤖 Generating AI summary..."):
                    summary_points = summarize_file(uploaded_file)
                
                st.markdown("""
                <div class="success-message">
                    ✅ AI summary generated with key insights
                </div>
                """, unsafe_allow_html=True)

                # Step 3: Map summary to paragraphs
                with st.spinner("🎯 Mapping insights to paragraphs..."):
                    mapping = map_summary_to_paragraphs(paragraphs, summary_points)
                
                st.markdown("""
                <div class="success-message">
                    ✅ Intelligent mapping completed successfully
                </div>
                """, unsafe_allow_html=True)
                
                # Store results and transition
                st.session_state.paragraphs = paragraphs
                st.session_state.mapping = mapping
                st.session_state.processing_complete = True
                
                # Final success message
                st.markdown("""
                <div style="background: linear-gradient(135deg, #00b894, #00a085); color: white; 
                     padding: 30px; border-radius: 20px; text-align: center; margin: 30px 0;
                     box-shadow: 0 15px 35px rgba(0, 184, 148, 0.4); animation: successPulse 1s ease-out;">
                    <div style="font-size: 3rem; margin-bottom: 15px;">🎉</div>
                    <h2 style="margin: 0 0 10px 0;">Analysis Complete!</h2>
                    <p style="margin: 0; opacity: 0.9; font-size: 1.1rem;">
                        Loading your beautiful document viewer...
                    </p>
                </div>
                """, unsafe_allow_html=True)
                
                st.rerun()

    # Show the clean document viewer interface
    if st.session_state.processing_complete:
        # Add a header with file info and reset option
        col1, col2 = st.columns([3, 1])
        with col1:
            st.markdown(f"""
            <div style="background: rgba(255,255,255,0.9); backdrop-filter: blur(15px); 
                 padding: 20px 30px; border-radius: 20px; margin: 20px 0;
                 box-shadow: 0 10px 30px rgba(0,0,0,0.1);">
                <div style="display: flex; align-items: center;">
                    <div style="font-size: 2rem; margin-right: 15px;">📄</div>
                    <div>
                        <h3 style="margin: 0; color: #2d3436; font-weight: 600;">
                            {uploaded_file.name}
                        </h3>
                        <p style="margin: 5px 0 0 0; color: #636e72;">
                            Analysis completed • Ready for review
                        </p>
                    </div>
                </div>
            </div>
            """, unsafe_allow_html=True)
        with col2:
            if st.button("🔄 New Document", type="secondary", key="reset_btn"):
                # Reset session state
                st.session_state.processing_complete = False
                st.session_state.paragraphs = None
                st.session_state.mapping = None
                st.rerun()

        # Get data from session state
        paragraphs = st.session_state.paragraphs
        mapping = st.session_state.mapping

        # -------------------- Clean Document Viewer --------------------
        # Initialize session state for highlighting
        if "highlighted_paragraph" not in st.session_state:
            st.session_state.highlighted_paragraph = None
        
        # Add modern, attractive CSS with glassmorphism and animations
        st.markdown("""
        <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
        
        * {
            font-family: 'Inter', sans-serif;
        }
        
        .main-container {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 50%, #f093fb 100%);
            min-height: 100vh;
            padding: 20px 0;
        }
        
        .document-viewer {
            background: rgba(255, 255, 255, 0.95);
            backdrop-filter: blur(20px);
            border-radius: 20px;
            padding: 30px;
            margin: 20px 0;
            box-shadow: 0 20px 40px rgba(0,0,0,0.1);
            border: 1px solid rgba(255,255,255,0.2);
        }
        
        .paragraph-box {
            background: rgba(255, 255, 255, 0.9);
            backdrop-filter: blur(15px);
            color: #2d3436;
            padding: 20px;
            margin: 15px 0;
            border-radius: 15px;
            border-left: 5px solid #6c5ce7;
            box-shadow: 0 8px 32px rgba(108, 92, 231, 0.1);
            line-height: 1.8;
            transition: all 0.4s cubic-bezier(0.4, 0.0, 0.2, 1);
            position: relative;
            overflow: hidden;
        }
        
        .paragraph-box::before {
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            height: 2px;
            background: linear-gradient(90deg, #6c5ce7, #a29bfe, #fd79a8);
            transform: scaleX(0);
            transition: transform 0.4s ease;
        }
        
        .paragraph-box:hover {
            transform: translateY(-5px);
            box-shadow: 0 15px 50px rgba(108, 92, 231, 0.2);
        }
        
        .paragraph-box:hover::before {
            transform: scaleX(1);
        }
        
        .summary-box {
            background: linear-gradient(135deg, #6c5ce7 0%, #a29bfe 50%, #fd79a8 100%);
            color: white;
            padding: 25px;
            margin: 20px 0;
            border-radius: 20px;
            box-shadow: 0 15px 35px rgba(108, 92, 231, 0.3);
            line-height: 1.8;
            position: relative;
            overflow: hidden;
            transition: all 0.4s cubic-bezier(0.4, 0.0, 0.2, 1);
        }
        
        .summary-box::before {
            content: '';
            position: absolute;
            top: -50%;
            left: -50%;
            width: 200%;
            height: 200%;
            background: linear-gradient(45deg, transparent, rgba(255,255,255,0.1), transparent);
            transform: rotate(45deg);
            transition: all 0.6s;
            opacity: 0;
        }
        
        .summary-box:hover {
            transform: translateY(-8px) scale(1.02);
            box-shadow: 0 25px 50px rgba(108, 92, 231, 0.4);
        }
        
        .summary-box:hover::before {
            opacity: 1;
            top: -100%;
            left: -100%;
        }
        
        .viewer-header {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 50%, #f093fb 100%);
            color: white;
            padding: 40px;
            border-radius: 25px;
            text-align: center;
            margin-bottom: 30px;
            position: relative;
            overflow: hidden;
            box-shadow: 0 20px 40px rgba(102, 126, 234, 0.3);
        }
        
        .viewer-header::before {
            content: '';
            position: absolute;
            top: 0;
            left: -100%;
            width: 100%;
            height: 100%;
            background: linear-gradient(90deg, transparent, rgba(255,255,255,0.2), transparent);
            animation: shimmer 3s infinite;
        }
        
        @keyframes shimmer {
            0% { left: -100%; }
            100% { left: 100%; }
        }
        
        .viewer-header h2 {
            margin: 0;
            font-weight: 700;
            font-size: 2.5em;
            text-shadow: 0 2px 10px rgba(0,0,0,0.2);
            animation: fadeInUp 1s ease-out;
        }
        
        .viewer-header p {
            margin: 10px 0 0 0;
            opacity: 0.9;
            font-size: 1.2em;
            animation: fadeInUp 1s ease-out 0.2s both;
        }
        
        @keyframes fadeInUp {
            from {
                opacity: 0;
                transform: translateY(30px);
            }
            to {
                opacity: 1;
                transform: translateY(0);
            }
        }
        
        .stats-card {
            background: rgba(255, 255, 255, 0.9);
            backdrop-filter: blur(20px);
            padding: 25px 20px;
            border-radius: 18px;
            box-shadow: 0 10px 30px rgba(0,0,0,0.1);
            text-align: center;
            margin: 15px 0;
            border: 1px solid rgba(255,255,255,0.3);
            transition: all 0.4s cubic-bezier(0.4, 0.0, 0.2, 1);
            position: relative;
            overflow: hidden;
        }
        
        .stats-card::before {
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            height: 3px;
            background: linear-gradient(90deg, #6c5ce7, #a29bfe, #fd79a8, #fdcb6e);
            transform: scaleX(0);
            transition: transform 0.6s ease;
        }
        
        .stats-card:hover {
            transform: translateY(-10px);
            box-shadow: 0 20px 40px rgba(0,0,0,0.15);
        }
        
        .stats-card:hover::before {
            transform: scaleX(1);
        }
        
        .stats-card h3 {
            background: linear-gradient(135deg, #6c5ce7, #a29bfe);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            font-weight: 700;
            font-size: 2.2em;
            margin: 0;
            animation: countUp 2s ease-out;
        }
        
        @keyframes countUp {
            from { transform: scale(0); }
            to { transform: scale(1); }
        }
        
        .section-header {
            background: linear-gradient(135deg, #2d3436, #636e72);
            color: white;
            padding: 20px;
            border-radius: 15px;
            margin: 20px 0 10px 0;
            text-align: center;
            box-shadow: 0 8px 25px rgba(45, 52, 54, 0.2);
            position: relative;
            overflow: hidden;
        }
        
        .section-header::after {
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            bottom: 0;
            background: linear-gradient(45deg, transparent, rgba(255,255,255,0.05), transparent);
            transform: translateX(-100%);
            transition: transform 0.8s;
        }
        
        .section-header:hover::after {
            transform: translateX(100%);
        }
        
        .floating-element {
            animation: float 6s ease-in-out infinite;
        }
        
        @keyframes float {
            0%, 100% { transform: translateY(0px); }
            50% { transform: translateY(-10px); }
        }
        
        .pulse-dot {
            display: inline-block;
            width: 8px;
            height: 8px;
            border-radius: 50%;
            background: #6c5ce7;
            animation: pulse 2s infinite;
            margin-right: 8px;
        }
        
        @keyframes pulse {
            0% { transform: scale(1); opacity: 1; }
            50% { transform: scale(1.2); opacity: 0.7; }
            100% { transform: scale(1); opacity: 1; }
        }
        
        .glow-text {
            text-shadow: 0 0 20px rgba(108, 92, 231, 0.5);
        }
        
        .scroll-indicator {
            position: fixed;
            top: 0;
            left: 0;
            height: 4px;
            background: linear-gradient(90deg, #6c5ce7, #a29bfe, #fd79a8);
            z-index: 1000;
            transition: width 0.1s;
        }
        </style>
        """, unsafe_allow_html=True)
        
        # Add document stats with enhanced animations
        st.markdown("""
        <div class="viewer-header floating-element">
            <h2 class="glow-text">🚀 Document Analysis Complete</h2>
            <p>Your legal document has been processed with AI-powered intelligence</p>
        </div>
        """, unsafe_allow_html=True)
        
        # Enhanced stats cards with icons and animations
        col1, col2, col3 = st.columns(3)
        with col1:
            st.markdown(f"""
            <div class="stats-card floating-element" style="animation-delay: 0.1s;">
                <div style="font-size: 2.5em; margin-bottom: 10px;">📄</div>
                <h3>{len(paragraphs)}</h3>
                <p style="color: #636e72; margin: 5px 0 0 0; font-weight: 500;">Paragraphs Analyzed</p>
            </div>
            """, unsafe_allow_html=True)
        
        with col2:
            st.markdown(f"""
            <div class="stats-card floating-element" style="animation-delay: 0.2s;">
                <div style="font-size: 2.5em; margin-bottom: 10px;">✨</div>
                <h3>{len(mapping)}</h3>
                <p style="color: #636e72; margin: 5px 0 0 0; font-weight: 500;">Key Insights</p>
            </div>
            """, unsafe_allow_html=True)
        
        with col3:
            total_words = sum(len(p['text'].split()) for p in paragraphs)
            st.markdown(f"""
            <div class="stats-card floating-element" style="animation-delay: 0.3s;">
                <div style="font-size: 2.5em; margin-bottom: 10px;">📊</div>
                <h3>{total_words:,}</h3>
                <p style="color: #636e72; margin: 5px 0 0 0; font-weight: 500;">Words Processed</p>
            </div>
            """, unsafe_allow_html=True)
        # Main content area with enhanced styling
        st.markdown("<br><br>", unsafe_allow_html=True)
            
        left_col, right_col = st.columns([1.3, 1], gap="large")

        with left_col:
            st.markdown("""
            <div class="section-header">
                <h3 style="margin: 0;"><span class="pulse-dot"></span>📖 Document Content</h3>
            </div>
            """, unsafe_allow_html=True)
            
            # Document container with enhanced styling
            with st.container():
                for i, p in enumerate(paragraphs):
                    st.markdown(f"""
                    <div class="paragraph-box" style="animation-delay: {i * 0.1}s;">
                        <div style="display: flex; align-items: center; margin-bottom: 12px;">
                            <div style="background: linear-gradient(135deg, #6c5ce7, #a29bfe); color: white; 
                                        padding: 5px 12px; border-radius: 20px; font-size: 0.85em; 
                                        font-weight: 600; margin-right: 12px;">
                                § {p['id'] + 1}
                            </div>
                            <div style="color: #636e72; font-size: 0.9em;">
                                {len(p['text'].split())} words
                            </div>
                        </div>
                        <div style="line-height: 1.8; color: #2d3436;">
                            {p['text']}
                        </div>
                    </div>
                    """, unsafe_allow_html=True)

        with right_col:
            st.markdown("""
            <div class="section-header">
                <h3 style="margin: 0;"><span class="pulse-dot"></span>🎯 Executive Summary</h3>
            </div>
            """, unsafe_allow_html=True)
            
            # Summary container with enhanced styling
            with st.container():
                for i, m in enumerate(mapping):
                    st.markdown(f"""
                    <div class="summary-box" style="animation-delay: {i * 0.15}s;">
                        <div style="display: flex; align-items: center; margin-bottom: 15px;">
                            <div style="background: rgba(255,255,255,0.2); color: white; 
                                        padding: 8px 15px; border-radius: 25px; font-size: 0.9em; 
                                        font-weight: 600; margin-right: 15px;">
                                💡 Insight {i+1}
                            </div>
                            <div style="color: rgba(255,255,255,0.8); font-size: 0.85em;">
                                High Priority
                            </div>
                        </div>
                        <div style="line-height: 1.8; font-size: 1.05em;">
                            {m['summary_sentence']}
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    # Enhanced mapping info
                    st.markdown(f"""
                    <div style="text-align: center; margin: -10px 0 25px 0;">
                        <span style="
                            background: rgba(0,0,0,0.05); 
                            color: #2d3436; 
                           padding: 8px 18px; 
                            border-radius: 25px; 
                            font-size: 0.9em; 
                            font-weight: 600; 
                            border: 1px solid rgba(0,0,0,0.2);
                            box-shadow: 0 2px 6px rgba(0,0,0,0.1);
                            display: inline-flex;
                            align-items: center;
                            gap: 6px;
                        ">
                            📍 References Paragraph {m['reference_paragraph_id'] + 1}
                        </span>
                    </div>

                    """, unsafe_allow_html=True)
