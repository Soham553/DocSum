import streamlit as st
from PyPDF2 import PdfReader
import docx
from sentence_transformers import SentenceTransformer, util
from google import genai
from google.genai import types
import io

# -------------------- Config --------------------
API_KEY = "AIzaSyAd-pJNXnxFNCr-zrb967WAILMuD6O-QxQ"
st.set_page_config(page_title="Legal Document Summarizer", layout="wide")

# Load embedding model
@st.cache_resource
def load_embedding_model():
    return SentenceTransformer("all-MiniLM-L6-v2")

model = load_embedding_model()

# -------------------- Helper Functions --------------------
def extract_paragraphs(uploaded_file):
    """
    FIXED: Now takes the uploaded_file object directly, not just the name
    """
    paragraphs = []
    
    # Get file extension from the name
    file_extension = uploaded_file.name.lower().split('.')[-1]

    if file_extension == "txt":
        # Read bytes and decode
        text = uploaded_file.read().decode("utf-8")
        paragraphs = text.split("\n\n")

    elif file_extension == "pdf":
        reader = PdfReader(io.BytesIO(uploaded_file.read()))
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        paragraphs = text.split("\n\n")

    elif file_extension == "docx":
        doc = docx.Document(io.BytesIO(uploaded_file.read()))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Assign IDs
    indexed_paragraphs = [{"id": idx, "text": para.strip()} for idx, para in enumerate(paragraphs) if para.strip()]
    return indexed_paragraphs


def summarize_file(uploaded_file, model_name="gemini-2.5-flash"):
    """
    FIXED: Now takes uploaded_file object and handles it properly
    """
    # Reset file pointer to beginning (important!)
    uploaded_file.seek(0)
    file_bytes = uploaded_file.read()
    
    # Determine MIME type based on file extension
    file_extension = uploaded_file.name.lower().split('.')[-1]
    if file_extension == 'pdf':
        mime_type = 'application/pdf'
    elif file_extension == 'docx':
        mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    else:  # txt
        mime_type = 'text/plain'

    file_part = types.Part.from_bytes(
        data=file_bytes,
        mime_type=mime_type,
    )

    prompt = "Please provide a clear and concise summary of the document in bullet points."

    client = genai.Client(api_key=API_KEY)
    response = client.models.generate_content(
        model=model_name,
        contents=[file_part, prompt]
    )

    # Split response into sentences/points
    summary_points = [p.strip() for p in response.text.split("\n") if p.strip()]
    return summary_points

def map_summary_to_paragraphs(paragraphs, summary_sentences):
    para_texts = [p["text"] for p in paragraphs]
    para_embeddings = model.encode(para_texts, convert_to_tensor=True)
    summary_embeddings = model.encode(summary_sentences, convert_to_tensor=True)

    mapping = []
    for idx, s_emb in enumerate(summary_embeddings):
        cos_scores = util.cos_sim(s_emb, para_embeddings)[0]
        best_match_idx = int(cos_scores.argmax())
        mapping.append({
            "summary_sentence": summary_sentences[idx],
            "reference_paragraph_id": paragraphs[best_match_idx]["id"],
            "reference_text": paragraphs[best_match_idx]["text"]
        })
    return mapping

# -------------------- Streamlit UI --------------------
st.title("📄 Legal Document Summarizer with Paragraph Mapping")

uploaded_file = st.file_uploader("Upload your legal document (.pdf, .docx, .txt)", type=["pdf","docx","txt"])

if uploaded_file:
    # Initialize processing state
    if "processing_complete" not in st.session_state:
        st.session_state.processing_complete = False
    if "paragraphs" not in st.session_state:
        st.session_state.paragraphs = None
    if "mapping" not in st.session_state:
        st.session_state.mapping = None

    # Show processing interface only if not complete
    if not st.session_state.processing_complete:
        with st.spinner("Extracting paragraphs..."):
            st.write(f"Processing: {uploaded_file.name}")
            paragraphs = extract_paragraphs(uploaded_file)
        
        st.success(f"✅ Extracted {len(paragraphs)} paragraphs.")

        if st.button("Generate Summary", type="primary"):
            with st.spinner("Generating summary..."):
                summary_points = summarize_file(uploaded_file)
            st.success("✅ Summary generated!")

            # Map summary to paragraphs
            with st.spinner("Mapping summary points to document paragraphs..."):
                mapping = map_summary_to_paragraphs(paragraphs, summary_points)
            st.success("✅ Mapping completed!")
            
            # Store results in session state
            st.session_state.paragraphs = paragraphs
            st.session_state.mapping = mapping
            st.session_state.processing_complete = True
            st.success("🎉 Processing complete! Loading document viewer...")
            st.rerun()

    # Show the clean document viewer interface
    if st.session_state.processing_complete:
        # Clear the page and show new interface
        st.markdown("---")
        
        # Add a header with file info and reset option
        col1, col2 = st.columns([3, 1])
        with col1:
            st.markdown(f"### 📄 {uploaded_file.name}")
        with col2:
            if st.button("🔄 Process New Document", type="secondary"):
                # Reset session state
                st.session_state.processing_complete = False
                st.session_state.paragraphs = None
                st.session_state.mapping = None
                st.rerun()

        # Get data from session state
        paragraphs = st.session_state.paragraphs
        mapping = st.session_state.mapping

        # -------------------- Clean Document Viewer --------------------
        # Initialize session state for highlighting
        if "highlighted_paragraph" not in st.session_state:
            st.session_state.highlighted_paragraph = None
        
        # Add custom CSS for the clean viewer interface
        st.markdown("""
        <style>
        .document-viewer {
            background: #f8f9fa;
            border-radius: 12px;
            padding: 20px;
            margin: 10px 0;
        }
        
        .paragraph-box {
            background: linear-gradient(135deg, #2c3e50 0%, #34495e 100%);
            color: white;
            padding: 15px;
            margin: 10px 0;
            border-radius: 8px;
            border-left: 4px solid #74b9ff;
            box-shadow: 0 1px 3px rgba(0,0,0,0.3);
            line-height: 1.6;
            transition: all 0.3s ease;
        }
        
        .summary-box {
            background: linear-gradient(135deg, #00b894 0%, #00a085 100%);
            color: white;
            padding: 15px;
            margin: 10px 0;
            border-radius: 8px;
            border-left: 4px solid #fff;
            box-shadow: 0 2px 6px rgba(0,184,148,0.3);
            line-height: 1.6;
        }
        
        .viewer-header {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 20px;
            border-radius: 12px 12px 0 0;
            text-align: center;
            margin-bottom: 0;
        }
        
        .stats-card {
            background: white;
            padding: 15px;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            text-align: center;
            margin: 10px 0;
        }
        </style>
        """, unsafe_allow_html=True)
        
        # Add document stats
        st.markdown("""
        <div class="viewer-header">
            <h2>📄 Document Analysis Complete</h2>
            <p>Your legal document has been processed and summarized</p>
        </div>
        """, unsafe_allow_html=True)
        
        # Stats cards
        col1, col2, col3 = st.columns(3)
        with col1:
            st.markdown(f"""
            <div class="stats-card">
                <h3 style="color: #2d3436; margin: 0;">{len(paragraphs)}</h3>
                <p style="color: #636e72; margin: 5px 0 0 0;">Paragraphs</p>
            </div>
            """, unsafe_allow_html=True)
        
        with col2:
            st.markdown(f"""
            <div class="stats-card">
                <h3 style="color: #2d3436; margin: 0;">{len(mapping)}</h3>
                <p style="color: #636e72; margin: 5px 0 0 0;">Summary Points</p>
            </div>
            """, unsafe_allow_html=True)
        
        with col3:
            total_words = sum(len(p['text'].split()) for p in paragraphs)
            st.markdown(f"""
            <div class="stats-card">
                <h3 style="color: #2d3436; margin: 0;">{total_words:,}</h3>
                <p style="color: #636e72; margin: 5px 0 0 0;">Total Words</p>
            </div>
            """, unsafe_allow_html=True)
        # Main content area with better spacing
        st.markdown("<br>", unsafe_allow_html=True)
            
        left_col, right_col = st.columns([1.3, 1], gap="large")

        with left_col:
            st.markdown("### 📑 Document Content")
            
            # Document container
            with st.container():
                for p in paragraphs:
                    st.markdown(f"""
                    <div class="paragraph-box" id="paragraph-{p['id']}">
                        <strong>Paragraph {p['id'] + 1}:</strong><br>
                        {p['text']}
                    </div>
                    """, unsafe_allow_html=True)

        with right_col:
            st.markdown("### 📝 Executive Summary")
            
            # Summary container
            with st.container():
                for i, m in enumerate(mapping):
                    st.markdown(f"""
                    <div class="summary-box">
                        <strong>📌 Key Point {i+1}:</strong><br>
                        {m['summary_sentence']}
                    </div>
                    """, unsafe_allow_html=True)
                    
                    # Show mapping info subtly
                    st.caption(f"📍 References paragraph {m['reference_paragraph_id'] + 1}")
                    st.markdown("<br>", unsafe_allow_html=True)
